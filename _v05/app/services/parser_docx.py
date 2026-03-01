"""Parsing and writing docx questionnaires with python-docx."""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from app.models.project import DocumentStructure
from app.services.anonymizer import Anonymizer

logger = logging.getLogger(__name__)


def extract_raw_content(filepath: Path, max_paragraphs: int = 50) -> str:
    """Extract first max_paragraphs paragraphs/table rows as text for Claude analysis."""
    doc = DocxDocument(str(filepath))
    lines = []
    count = 0

    for element in doc.element.body:
        if count >= max_paragraphs:
            break
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if text:
                lines.append(f"[{style}] {text}")
                count += 1
        elif tag == "tbl":
            from docx.table import Table
            table = Table(element, doc)
            for i, row in enumerate(table.rows[:8]):
                cells = " | ".join(c.text.strip() for c in row.cells)
                if cells.strip():
                    lines.append(f"[Table ligne {i + 1}] {cells}")
                    count += 1

    return "\n".join(lines)


def extract_questions(filepath: Path, structure: DocumentStructure) -> list[dict]:
    """
    Extract questions from docx using the response marker.
    Returns list of {id, question, paragraph_index}.
    """
    response_marker = (structure.response_marker or "Réponse du titulaire").lower()
    doc = DocxDocument(str(filepath))
    paragraphs = list(doc.paragraphs)
    questions = []
    q_counter = 0

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if response_marker in text.lower():
            # Search backwards for the question text (skip empty + marker lines)
            q_text = ""
            for j in range(max(0, i - 6), i):
                candidate = paragraphs[j].text.strip()
                if candidate and response_marker not in candidate.lower():
                    q_text = candidate

            q_counter += 1
            q_id = f"Q{q_counter:03d}"
            if q_text:
                questions.append({
                    "id": q_id,
                    "question": q_text,
                    "paragraph_index": i,
                })

    # Also scan tables for response markers
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if response_marker in para.text.strip().lower():
                        q_counter += 1
                        q_id = f"Q{q_counter:03d}"
                        questions.append({
                            "id": q_id,
                            "question": para.text.strip(),
                            "in_table": True,
                        })

    return questions


def format_questionnaire_for_claude(questions: list[dict]) -> str:
    """Format extracted questions as text block for the Claude prompt."""
    lines = []
    for q in questions:
        lines.append(f"ID: {q['id']}")
        lines.append(f"Question: {q['question']}")
        lines.append("")
    return "\n".join(lines)


def write_responses(
    filepath: Path, responses: list[dict], structure: DocumentStructure
) -> None:
    """Write Claude responses into response zones of the docx file."""
    response_map = {r["question_id"]: r["response"] for r in responses}
    response_marker = (structure.response_marker or "Réponse du titulaire").lower()

    doc = DocxDocument(str(filepath))
    paragraphs = list(doc.paragraphs)
    q_counter = 0

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if response_marker in text.lower():
            q_counter += 1
            q_id = f"Q{q_counter:03d}"
            if q_id in response_map:
                # Write into the next paragraph if it's empty, else append a run
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    if not next_para.text.strip():
                        if next_para.runs:
                            next_para.runs[0].text = response_map[q_id]
                        else:
                            next_para.add_run(response_map[q_id])
                    else:
                        para.add_run("\n" + response_map[q_id])
                else:
                    para.add_run("\n" + response_map[q_id])

    doc.save(str(filepath))


def apply_anonymization(filepath: Path, anonymizer: Anonymizer) -> None:
    """Apply anonymization to all text runs in the docx."""
    doc = DocxDocument(str(filepath))
    _apply_to_doc(doc, anonymizer.anonymize)
    doc.save(str(filepath))


def apply_deanonymization(filepath: Path, anonymizer: Anonymizer) -> None:
    """Apply deanonymization to all text runs in the docx."""
    doc = DocxDocument(str(filepath))
    _apply_to_doc(doc, anonymizer.deanonymize)
    doc.save(str(filepath))


def _apply_to_doc(doc, transform) -> None:
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = transform(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = transform(run.text)
