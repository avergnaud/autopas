"""PAS Assistant — Response generation pipeline.

Orchestrates the full generation workflow:
  1. Load corpus references
  2. Read questions from anonymized xlsx
  3. Call Claude to generate responses (JSON)
  4. Write responses into output xlsx
  5. Call Claude to generate attention points (JSON)
  6. De-anonymize output xlsx and attention text
  7. Save attention.md
"""

import json
import logging
import os
import re
from pathlib import Path

import anthropic
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string

from app.config import BASE_DIR, get_config
from app.services import project_manager
from app.services.anonymizer import deanonymize_text, deanonymize_xlsx
from app.services.parser_xlsx import read_questions, read_status_choices, write_responses

logger = logging.getLogger(__name__)

PROJECTS_DIR = BASE_DIR / "data" / "projects"

# ---------------------------------------------------------------------------
# French article correction (traitement B — avant dé-anonymisation)
# ---------------------------------------------------------------------------

_TOKENS = r'(?:FOURNISSEUR|CLIENT)'

_FRENCH_ARTICLE_FIXES = [
    (re.compile(r'\bau\s+(' + _TOKENS + r')\b'),    r'à \1'),
    (re.compile(r'\bdu\s+(' + _TOKENS + r')\b'),    r'de \1'),
    (re.compile(r'\b[Ll]e\s+(' + _TOKENS + r')\b'), r'\1'),
    (re.compile(r'\b[Ll]a\s+(' + _TOKENS + r')\b'), r'\1'),
]


def fix_french_token_articles(text: str) -> str:
    """Corrige les articles contractés/définis devant FOURNISSEUR/CLIENT (avant dé-anonymisation)."""
    for pattern, repl in _FRENCH_ARTICLE_FIXES:
        text = pattern.sub(repl, text)
    return text
CORPUS_DIR = BASE_DIR / "data" / "corpus"
PROMPTS_DIR = BASE_DIR / "data" / "config" / "prompts"

_VERBOSITY_MAP = {"Concis": 1, "Standard": 2, "Détaillé": 3}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _load_prompt(name: str) -> str:
    """Load a prompt file from data/config/prompts/."""
    path = PROMPTS_DIR / name
    if not path.exists():
        raise RuntimeError(f"Fichier prompt introuvable : {path}")
    return path.read_text(encoding="utf-8")


def _resolve_verbosity(cadrage: dict, config: dict) -> dict:
    """Return verbosity dict {level, label, max_words} from cadrage answer.

    Args:
        cadrage: Cadrage answers dict (may contain 'verbosity' key).
        config: Full application config.

    Returns:
        Dict with level (int), label (str), max_words (int).
    """
    label = cadrage.get("verbosity", "Standard")
    level = _VERBOSITY_MAP.get(label, 2)
    levels_config = config.get("verbosity", {}).get("levels", {})
    level_cfg = levels_config.get(level, levels_config.get(str(level), {}))
    return {
        "level": level,
        "label": level_cfg.get("label", label),
        "max_words": level_cfg.get("max_words", 100),
    }


def _format_cadrage(cadrage: dict) -> str:
    """Format cadrage answers as human-readable text for Claude.

    Args:
        cadrage: Cadrage answers dict.

    Returns:
        Multi-line string with key: value pairs, excluding 'verbosity'.
    """
    labels = {
        "pas_niveau_entreprise": "PAS niveau entreprise",
        "type_prestation_base": "Type de prestation",
        "type_prestation_detail": "Détail type de prestation",
        "nb_etp": "Nombre d'ETP",
        "activites": "Activités de la prestation",
        "expertise_atlassian": "Expertise Atlassian",
        "hebergement_donnees": "Hébergement des données",
        "cloud_provider": "Fournisseur Cloud",
        "sous_traitance_rgpd": "Sous-traitance RGPD",
        "lieu_travail": "Lieu de travail",
        "agences": "Agences concernées",
        "poste_travail": "Type de poste de travail",
        "connexion_distante": "Mode de connexion distante",
        "secteur_client": "Secteur du CLIENT",
    }
    lines = []
    for key, label in labels.items():
        val = cadrage.get(key)
        if val is None or val == "":
            continue
        if isinstance(val, list):
            val = ", ".join(str(v) for v in val)
        lines.append(f"- {label} : {val}")
    return "\n".join(lines) if lines else "(aucun contexte renseigné)"


def _read_corpus_entry_xlsx(corpus_dir: Path, meta: dict) -> str:
    """Extract Q&A content from a corpus xlsx file.

    Args:
        corpus_dir: Path to the corpus entry directory.
        meta: Metadata dict for this corpus entry.

    Returns:
        Formatted text content of the corpus file.
    """
    anon_path = corpus_dir / "anonymized.xlsx"
    if not anon_path.exists():
        return ""

    struct_path = corpus_dir / "structure.json"
    structure = json.loads(struct_path.read_text(encoding="utf-8")) if struct_path.exists() else {}

    wb = load_workbook(anon_path, read_only=True, data_only=True)

    sheet_name = structure.get("selected_sheet")
    if sheet_name and sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        ws = wb.active

    lines: list[str] = [f"=== {meta.get('filename', corpus_dir.name)} ==="]

    if structure.get("col_question"):
        first_data_row = int(structure.get("first_data_row") or 2)
        col_q_idx = column_index_from_string(structure["col_question"].upper())
        col_r_idx = (
            column_index_from_string(structure["col_response"].upper())
            if structure.get("col_response")
            else None
        )
        col_id_idx = (
            column_index_from_string(structure["col_id"].upper())
            if structure.get("col_id")
            else None
        )

        count = 0
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if row_idx < first_data_row:
                continue
            if count >= 150:
                break

            q_raw = row[col_q_idx - 1] if col_q_idx - 1 < len(row) else None
            q_text = str(q_raw).strip() if q_raw is not None else ""
            if not q_text:
                continue

            q_id = str(row_idx)
            if col_id_idx is not None and col_id_idx - 1 < len(row):
                id_raw = row[col_id_idx - 1]
                if id_raw is not None:
                    q_id = str(id_raw).strip()

            r_text = ""
            if col_r_idx is not None and col_r_idx - 1 < len(row):
                r_raw = row[col_r_idx - 1]
                r_text = str(r_raw).strip() if r_raw is not None else ""

            line = f"ID: {q_id} | Question: {q_text}"
            if r_text:
                line += f" | Réponse: {r_text}"
            lines.append(line)
            count += 1
    else:
        # No confirmed structure — dump non-empty cells row by row
        count = 0
        for row in ws.iter_rows(values_only=True):
            if count >= 200:
                break
            row_text = " | ".join(str(c).strip() for c in row if c)
            if row_text:
                lines.append(row_text)
                count += 1

    wb.close()
    return "\n".join(lines)


def _read_corpus_entry_docx(corpus_dir: Path, meta: dict) -> str:
    """Extract text content from a corpus docx file.

    Args:
        corpus_dir: Path to the corpus entry directory.
        meta: Metadata dict for this corpus entry.

    Returns:
        Formatted text content of the corpus file.
    """
    anon_path = corpus_dir / "anonymized.docx"
    if not anon_path.exists():
        return ""

    from docx import Document as DocxDocument  # local import to avoid hard dep at module level

    doc = DocxDocument(anon_path)
    parts: list[str] = [f"=== {meta.get('filename', corpus_dir.name)} ==="]

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    # Limit size
    if len(text) > 50_000:
        text = text[:50_000] + "\n[... tronqué]"
    return text


def _read_corpus_entry(corpus_id: str) -> str:
    """Read and format a single corpus entry for inclusion in Claude prompt.

    Args:
        corpus_id: UUID of the corpus entry.

    Returns:
        Formatted text content, or empty string if entry is unreadable.
    """
    corpus_dir = CORPUS_DIR / corpus_id
    meta_path = corpus_dir / "metadata.json"
    if not meta_path.exists():
        logger.warning("Corpus entry %s has no metadata.json", corpus_id)
        return ""

    meta = json.loads(meta_path.read_text(encoding="utf-8"))
    fmt = meta.get("format", "xlsx")

    try:
        if fmt == "xlsx":
            return _read_corpus_entry_xlsx(corpus_dir, meta)
        elif fmt == "docx":
            return _read_corpus_entry_docx(corpus_dir, meta)
        else:
            logger.warning("Corpus entry %s has unsupported format: %s", corpus_id, fmt)
            return ""
    except Exception:
        logger.exception("Failed to read corpus entry %s", corpus_id)
        return ""


def _build_user_prompt_responses(
    cadrage: dict,
    corpus_contents: list[str],
    questions: list[dict],
    verbosity: dict,
    status_choices: list[str] | None = None,
) -> str:
    """Build the user message for Claude's response-generation call.

    Args:
        cadrage: Cadrage answers dict.
        corpus_contents: List of formatted corpus text strings.
        questions: List of {"question_id", "question_text"} dicts.
        verbosity: Dict with label and max_words.
        status_choices: Optional list of allowed status values from the dropdown.

    Returns:
        Complete user prompt string.
    """
    lines: list[str] = []

    lines.append("=== CONTEXTE DE LA PRESTATION ===")
    lines.append(_format_cadrage(cadrage))
    lines.append("")

    if corpus_contents:
        lines.append("=== EXEMPLES DE QUESTIONNAIRES DÉJÀ REMPLIS ===")
        lines.append("\n\n".join(c for c in corpus_contents if c))
        lines.append("")

    if status_choices:
        lines.append("=== VALEURS DE STATUT AUTORISÉES ===")
        lines.append(
            "Pour chaque réponse, renseigner le champ \"status\" avec EXACTEMENT "
            "l'une de ces valeurs (respecter la casse) :"
        )
        for choice in status_choices:
            lines.append(f"- {choice}")
        lines.append("")

    lines.append("=== QUESTIONNAIRE À REMPLIR ===")
    lines.append(
        f"Niveau de verbosité : {verbosity['label']} — "
        f"maximum {verbosity['max_words']} mots par réponse."
    )
    lines.append("")

    for q in questions:
        lines.append(f"ID: {q['question_id']}")
        lines.append(f"Question: {q['question_text']}")
        lines.append("---")

    return "\n".join(lines)


def _build_user_prompt_attention(
    cadrage: dict,
    questions: list[dict],
    responses: list[dict],
) -> str:
    """Build the user message for Claude's attention-point generation call.

    Args:
        cadrage: Cadrage answers dict.
        questions: List of {"question_id", "question_text"} dicts.
        responses: List of {"question_id", "response"} dicts.

    Returns:
        Complete user prompt string.
    """
    resp_map = {r["question_id"]: r["response"] for r in responses}

    lines: list[str] = []
    lines.append("=== CONTEXTE DE LA PRESTATION ===")
    lines.append(_format_cadrage(cadrage))
    lines.append("")

    lines.append("=== QUESTIONNAIRE REMPLI ===")
    for q in questions:
        qid = q["question_id"]
        lines.append(f"ID: {qid}")
        lines.append(f"Question: {q['question_text']}")
        r = resp_map.get(qid, "(non répondu)")
        lines.append(f"Réponse: {r}")
        lines.append("---")

    return "\n".join(lines)


def _call_claude_json(system_prompt: str, user_prompt: str, model: str, max_tokens: int) -> dict:
    """Call Claude and parse the JSON response.

    Args:
        system_prompt: System prompt string.
        user_prompt: User message string.
        model: Claude model ID.
        max_tokens: Maximum tokens for the response.

    Returns:
        Parsed JSON dict from Claude's response.

    Raises:
        RuntimeError: If the API key is missing or the response is invalid JSON.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or api_key.startswith("sk-ant-..."):
        raise RuntimeError("ANTHROPIC_API_KEY non configurée.")

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )

    raw = response.content[0].text.strip()

    # Strip markdown code block if present
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.lower().startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)


def _format_attention_markdown(
    attention_points: list[dict],
    mapping: dict[str, str],
) -> str:
    """Format attention points as a Markdown document with de-anonymization.

    Args:
        attention_points: List of attention point dicts from Claude.
        mapping: Original anonymization mapping {original: token} for de-anonymization.

    Returns:
        Markdown string.
    """
    lines: list[str] = ["# Points d'attention\n"]

    for i, p in enumerate(attention_points, 1):
        qid = deanonymize_text(str(p.get("question_id", "")), mapping)
        cat = str(p.get("category", "INFORMATION"))
        desc = deanonymize_text(str(p.get("description", "")), mapping)
        rec = deanonymize_text(str(p.get("recommendation", "")), mapping)

        lines.append(f"## {i}. [{cat}] {qid}\n")
        lines.append(f"{desc}\n")
        lines.append(f"**Recommandation :** {rec}\n")
        lines.append("---\n")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main generation pipeline
# ---------------------------------------------------------------------------


def run_generation(project_id: str) -> None:
    """Entry point for the background generation task.

    Wraps _do_generation with error handling: on any exception, the project
    status is set to 'error' with the exception message.

    Args:
        project_id: UUID of the project to generate.
    """
    try:
        _do_generation(project_id)
    except Exception as exc:
        logger.exception("Generation failed for project %s", project_id)
        try:
            project_manager.update_project(
                project_id,
                status="error",
                error_message=str(exc),
            )
        except Exception:
            pass


def _do_generation(project_id: str) -> None:
    """Full generation pipeline.

    Args:
        project_id: UUID of the project.

    Raises:
        Various exceptions on failure (caught by run_generation).
    """
    project_dir = PROJECTS_DIR / project_id

    # -------------------------------------------------------------------------
    # Step 1 — Load project data
    # -------------------------------------------------------------------------
    proj = project_manager.load_project(project_id)

    cadrage = proj.get("cadrage")
    if not cadrage:
        raise RuntimeError("Cadrage non effectué.")

    selected_corpus: list[str] = proj.get("selected_corpus") or []

    # Load structure from separate file (more reliable than project.json copy)
    structure_path = project_dir / "structure.json"
    if not structure_path.exists():
        raise RuntimeError("Structure xlsx non confirmée.")
    structure = json.loads(structure_path.read_text(encoding="utf-8"))

    if not structure.get("col_question") or not structure.get("col_response"):
        raise RuntimeError(
            "Structure incomplète : col_question et col_response sont obligatoires."
        )

    # -------------------------------------------------------------------------
    # Step 2 — Init
    # -------------------------------------------------------------------------
    project_manager.update_project(
        project_id, status="generating", progress_step="Chargement des références..."
    )

    config = get_config()
    model = os.environ.get("CLAUDE_MODEL", config.get("claude", {}).get("model", "claude-sonnet-4-6"))
    max_tokens = int(config.get("claude", {}).get("max_tokens", 16000))
    max_files = int(config.get("reference", {}).get("max_files", 3))

    verbosity = _resolve_verbosity(cadrage, config)
    logger.info(
        "Project %s — model=%s verbosity=%s max_tokens=%d",
        project_id, model, verbosity["label"], max_tokens,
    )

    # -------------------------------------------------------------------------
    # Step 3 — Load corpus content
    # -------------------------------------------------------------------------
    corpus_ids_to_use = selected_corpus[:max_files]
    corpus_contents: list[str] = []
    for cid in corpus_ids_to_use:
        content = _read_corpus_entry(cid)
        if content:
            corpus_contents.append(content)
    logger.info("Loaded %d corpus entries", len(corpus_contents))

    # -------------------------------------------------------------------------
    # Step 4 — Read questions
    # -------------------------------------------------------------------------
    project_manager.update_project(project_id, progress_step="Lecture du questionnaire...")

    anonymized_path = project_dir / "anonymized.xlsx"
    if not anonymized_path.exists():
        raise RuntimeError("Fichier anonymisé introuvable.")

    questions = read_questions(anonymized_path, structure)
    if not questions:
        raise RuntimeError("Aucune question extraite du questionnaire.")

    logger.info("Extracted %d questions from project %s", len(questions), project_id)

    # Read status dropdown choices if a status column is defined
    status_choices = read_status_choices(anonymized_path, structure)
    if status_choices:
        logger.info("Status choices for project %s: %s", project_id, status_choices)
    else:
        logger.info("No status choices found for project %s (col_status=%s)", project_id, structure.get("col_status"))

    # -------------------------------------------------------------------------
    # Step 5 — Generate responses via Claude
    # -------------------------------------------------------------------------
    project_manager.update_project(
        project_id, progress_step="Génération des réponses (appel Claude)..."
    )

    system_response = _load_prompt("system_response.txt")
    user_prompt_resp = _build_user_prompt_responses(
        cadrage, corpus_contents, questions, verbosity, status_choices
    )

    logger.info(
        "Calling Claude for responses — prompt length: ~%d chars",
        len(system_response) + len(user_prompt_resp),
    )
    result_responses = _call_claude_json(system_response, user_prompt_resp, model, max_tokens)

    responses: list[dict] = result_responses.get("responses", [])
    if not responses:
        raise RuntimeError("Claude n'a retourné aucune réponse.")

    logger.info("Claude returned %d responses", len(responses))

    # Apply French article correction before de-anonymization (traitement B)
    for r in responses:
        if "response" in r:
            r["response"] = fix_french_token_articles(r["response"])

    # -------------------------------------------------------------------------
    # Step 6 — Write responses into output_anon.xlsx
    # -------------------------------------------------------------------------
    project_manager.update_project(project_id, progress_step="Écriture des réponses...")

    output_anon_path = project_dir / "output_anon.xlsx"
    write_responses(anonymized_path, output_anon_path, structure, responses)

    # -------------------------------------------------------------------------
    # Step 7 — Generate attention points via Claude
    # -------------------------------------------------------------------------
    project_manager.update_project(
        project_id, progress_step="Génération des points d'attention..."
    )

    system_attention = _load_prompt("system_attention.txt")
    user_prompt_attn = _build_user_prompt_attention(cadrage, questions, responses)

    logger.info(
        "Calling Claude for attention points — prompt length: ~%d chars",
        len(system_attention) + len(user_prompt_attn),
    )
    result_attention = _call_claude_json(system_attention, user_prompt_attn, model, max_tokens)

    attention_points: list[dict] = result_attention.get("attention_points", [])
    logger.info("Claude returned %d attention points", len(attention_points))

    # Apply French article correction on attention points before de-anonymization (traitement B)
    for p in attention_points:
        for field in ("description", "recommendation"):
            if field in p:
                p[field] = fix_french_token_articles(p[field])

    # -------------------------------------------------------------------------
    # Step 8 — De-anonymize output.xlsx
    # -------------------------------------------------------------------------
    project_manager.update_project(project_id, progress_step="Dé-anonymisation...")

    map_path = project_dir / "anonymized_map.json"
    if not map_path.exists():
        raise RuntimeError("Table de correspondance anonymization introuvable.")
    anon_mapping: dict[str, str] = json.loads(map_path.read_text(encoding="utf-8"))

    output_path = project_dir / "output.xlsx"
    deanonymize_xlsx(output_anon_path, output_path, anon_mapping)

    # Clean up intermediate file
    output_anon_path.unlink(missing_ok=True)

    # -------------------------------------------------------------------------
    # Step 9 — Format and save attention.md
    # -------------------------------------------------------------------------
    project_manager.update_project(project_id, progress_step="Finalisation...")

    attention_md = _format_attention_markdown(attention_points, anon_mapping)
    (project_dir / "attention.md").write_text(attention_md, encoding="utf-8")

    # -------------------------------------------------------------------------
    # Step 10 — Mark completed
    # -------------------------------------------------------------------------
    project_manager.update_project(
        project_id,
        status="completed",
        progress_step="Terminé.",
        error_message=None,
    )
    logger.info("Generation completed successfully for project %s", project_id)
