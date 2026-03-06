"""Anonymization utilities for xlsx and docx files.

Handles:
- Metadata extraction from docProps/core.xml and docProps/app.xml (xlsx)
  and from core_properties (docx)
- Keyword-based text replacement (sorted by length descending)
- File metadata stripping
- Preservation of local defined names (Excel dropdowns) across the roundtrip
"""

import logging
import re
import shutil
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from openpyxl.workbook.defined_name import DefinedName

logger = logging.getLogger(__name__)

_XLSX_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"

# Namespaces used in docProps/core.xml
_NS_DC = "http://purl.org/dc/elements/1.1/"
_NS_CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
_NS_DCTERMS = "http://purl.org/dc/terms/"

# Namespace used in docProps/app.xml
_NS_APP = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"

# Minimal clean app.xml to replace the original (strips company, app name, etc.)
_CLEAN_APP_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/'
    'extended-properties"></Properties>'
)


def safe_local_defined_names(xlsx_path: Path) -> list[tuple[str, str, int | None]]:
    """Extract safe, local defined names directly from workbook.xml.

    openpyxl's DefinedNameDict is keyed by name only, so when the same name
    appears twice (once sheet-scoped with localSheetId, once as a broken
    workbook-level external ref), the external #REF! entry silently overwrites
    the valid local one.  We read the raw XML first to capture local definitions
    that openpyxl would lose.

    Returns list of (name, attr_text, local_sheet_id) for names that:
    - reference cells inside this workbook (no [n] external-workbook notation)
    - are not broken (#REF!)
    - are not _xlnm.* built-ins
    - are not hidden
    """
    with zipfile.ZipFile(xlsx_path) as z:
        with z.open("xl/workbook.xml") as f:
            tree = ET.parse(f)

    results = []
    for dn in tree.iter(f"{{{_XLSX_NS}}}definedName"):
        name = dn.get("name", "")
        value = dn.text or ""
        hidden = dn.get("hidden", "0") == "1"
        local_sheet_id_raw = dn.get("localSheetId")

        if name.startswith("_xlnm."):
            continue
        if hidden:
            continue
        if "#REF" in value or re.search(r"\[\d+\]", value):
            continue

        sid = int(local_sheet_id_raw) if local_sheet_id_raw is not None else None
        results.append((name, value, sid))

    return results


def extract_metadata(xlsx_path: Path) -> dict[str, str]:
    """Extract PII-containing metadata fields from an xlsx file.

    Reads both docProps/core.xml (author, title, subject, …) and
    docProps/app.xml (company, application name) directly from the zip.

    Args:
        xlsx_path: Path to the xlsx file.

    Returns:
        Dict mapping field names to non-empty string values.
    """
    meta: dict[str, str] = {}

    with zipfile.ZipFile(xlsx_path) as z:
        names = z.namelist()

        # --- core.xml ---
        if "docProps/core.xml" in names:
            with z.open("docProps/core.xml") as f:
                tree = ET.parse(f)
            root = tree.getroot()

            def _core(tag_ns: str, tag_local: str) -> str:
                el = root.find(f"{{{tag_ns}}}{tag_local}")
                return (el.text or "").strip() if el is not None else ""

            _fields = [
                ("creator", _NS_DC, "creator"),
                ("lastModifiedBy", _NS_CP, "lastModifiedBy"),
                ("title", _NS_DC, "title"),
                ("subject", _NS_DC, "subject"),
                ("description", _NS_DC, "description"),
                ("keywords", _NS_CP, "keywords"),
            ]
            for key, ns, local in _fields:
                val = _core(ns, local)
                if val:
                    meta[key] = val

        # --- app.xml ---
        if "docProps/app.xml" in names:
            with z.open("docProps/app.xml") as f:
                tree = ET.parse(f)
            root = tree.getroot()

            for local in ("Company", "Application"):
                el = root.find(f"{{{_NS_APP}}}{local}")
                val = (el.text or "").strip() if el is not None else ""
                if val:
                    meta[local.lower()] = val

    logger.info("Extracted metadata from %s: %s", xlsx_path.name, list(meta.keys()))
    return meta


def anonymize_xlsx(
    source_path: Path,
    dest_path: Path,
    mapping: dict[str, str],
) -> dict[str, str]:
    """Replace keywords in all cell values and strip file metadata.

    Keywords are sorted by length descending to avoid partial replacements
    (e.g., "Ministère des Armées" is replaced before "Armées").

    Preserves local defined names (Excel validation dropdowns) using the same
    safe_local_defined_names technique as the openpyxl roundtrip.

    Args:
        source_path: Path to the source xlsx file (working.xlsx).
        dest_path: Path to save the anonymized xlsx file.
        mapping: Dict of {original_keyword: replacement} pairs.

    Returns:
        The mapping dict, sorted by key length descending (replacement order).
    """
    # Sort by key length descending to avoid partial replacements
    mapping = dict(
        sorted(
            {k.strip(): v.strip() for k, v in mapping.items() if k.strip()}.items(),
            key=lambda item: len(item[0]),
            reverse=True,
        )
    )

    # Capture safe local defined names before openpyxl touches the file
    safe_names = safe_local_defined_names(source_path)

    wb = load_workbook(source_path, keep_links=False)

    # --- Apply keyword replacements in all cell values ---
    replaced_count = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    original = cell.value
                    val = cell.value
                    for kw, token in mapping.items():
                        val = val.replace(kw, token)
                    if val != original:
                        cell.value = val
                        replaced_count += 1

    # --- Strip core metadata (openpyxl 3.x: wb.properties) ---
    cp = wb.properties
    cp.creator = "ANONYME"
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.description = ""
    cp.keywords = ""

    # --- Restore safe defined names (preserve Excel dropdowns) ---
    wb_dn_count = len(wb.defined_names)
    wb.defined_names.clear()
    ws_dn_count = 0
    for ws in wb.worksheets:
        ws_dn_count += len(ws.defined_names)
        ws.defined_names.clear()
        ws.print_area = None
        ws.print_title_rows = None
        ws.print_title_cols = None

    for name, attr_text, local_sheet_id in safe_names:
        dn = DefinedName(name=name, attr_text=attr_text, localSheetId=local_sheet_id)
        wb.defined_names.add(dn)

    logger.info(
        "Defined names: cleared %d wb-level + %d ws-level, re-injected %d safe: %s",
        wb_dn_count,
        ws_dn_count,
        len(safe_names),
        [n for n, _, _ in safe_names],
    )

    wb.save(dest_path)
    wb.close()

    # --- Patch app.xml in the saved zip to remove company / app info ---
    _strip_app_xml(dest_path)

    logger.info(
        "Anonymized %s -> %s: %d keywords, %d cells modified",
        source_path.name,
        dest_path.name,
        len(mapping),
        replaced_count,
    )

    return mapping


def extract_metadata_docx(docx_path: Path) -> dict[str, str]:
    """Extract PII-containing metadata fields from a docx file.

    Args:
        docx_path: Path to the docx file.

    Returns:
        Dict mapping field names to non-empty string values.
    """
    doc = DocxDocument(docx_path)
    cp = doc.core_properties
    meta: dict[str, str] = {}
    for key, val in [
        ("creator", cp.author),
        ("lastModifiedBy", cp.last_modified_by),
        ("title", cp.title),
        ("subject", cp.subject),
        ("keywords", cp.keywords),
        ("company", cp.company),
    ]:
        if val and str(val).strip():
            meta[key] = str(val).strip()
    logger.info("Extracted docx metadata from %s: %s", docx_path.name, list(meta.keys()))
    return meta


def anonymize_docx(
    source_path: Path,
    dest_path: Path,
    mapping: dict[str, str],
) -> dict[str, str]:
    """Replace keywords in all docx text (paragraphs + tables) and strip metadata.

    Keywords are sorted by length descending to avoid partial replacements.

    Args:
        source_path: Path to the source docx file.
        dest_path: Path to save the anonymized docx file.
        mapping: Dict of {original_keyword: replacement} pairs.

    Returns:
        The mapping dict, sorted by key length descending.
    """
    mapping = dict(
        sorted(
            {k.strip(): v.strip() for k, v in mapping.items() if k.strip()}.items(),
            key=lambda item: len(item[0]),
            reverse=True,
        )
    )

    shutil.copy2(source_path, dest_path)
    doc = DocxDocument(dest_path)

    def _replace_in_paragraph(para) -> int:
        count = 0
        for run in para.runs:
            if run.text:
                original = run.text
                for kw, token in mapping.items():
                    run.text = run.text.replace(kw, token)
                if run.text != original:
                    count += 1
        return count

    replaced_count = 0
    for para in doc.paragraphs:
        replaced_count += _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replaced_count += _replace_in_paragraph(para)

    cp = doc.core_properties
    cp.author = "ANONYME"
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.company = ""

    doc.save(dest_path)

    logger.info(
        "Anonymized docx %s -> %s: %d keywords, %d runs modified",
        source_path.name, dest_path.name, len(mapping), replaced_count,
    )
    return mapping


def _strip_app_xml(xlsx_path: Path) -> None:
    """Replace docProps/app.xml with a minimal clean version (in-place).

    Rewrites the xlsx zip, substituting a clean app.xml that contains no
    company name, application name, or other extended properties.
    """
    tmp_path = xlsx_path.with_suffix(".tmp.xlsx")
    try:
        with zipfile.ZipFile(xlsx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "docProps/app.xml":
                        zout.writestr(item, _CLEAN_APP_XML)
                    else:
                        zout.writestr(item, zin.read(item.filename))
        shutil.move(str(tmp_path), str(xlsx_path))
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise
