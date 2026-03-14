"""Anonymization utilities for xlsx and docx files.

Handles:
- Metadata extraction from docProps/core.xml and docProps/app.xml (xlsx)
  and from core_properties (docx)
- Keyword-based text replacement (sorted by length descending)
- File metadata stripping
- Preservation of local defined names (Excel dropdowns) across the roundtrip
"""

import html
import logging
import re
import shutil
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

_XLSX_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"

# Namespaces used in docProps/core.xml
_NS_DC = "http://purl.org/dc/elements/1.1/"
_NS_CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
_NS_DCTERMS = "http://purl.org/dc/terms/"

# Namespace used in docProps/app.xml
_NS_APP = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"

# Minimal clean app.xml to replace the original (strips company, app name, etc.)
_CLEAN_APP_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/'
    'extended-properties"></Properties>'
)

# Minimal clean core.xml to replace the original (strips all PII metadata fields)
_CLEAN_CORE_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
    '<cp:coreProperties'
    ' xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"'
    ' xmlns:dc="http://purl.org/dc/elements/1.1/"'
    ' xmlns:dcterms="http://purl.org/dc/terms/"'
    ' xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
    '<dc:creator>ANONYME</dc:creator>'
    '</cp:coreProperties>'
)

# XML files inside an xlsx zip that may contain user-visible text to anonymize
_XLSX_TEXT_TARGETS = frozenset(["xl/sharedStrings.xml"])

# Catamania — correction post-dé-anonymisation (case-insensitive)
_CATAMANIA_RE = re.compile(r'cat-amania', re.IGNORECASE)


def _is_worksheet_xml(fname: str) -> bool:
    """Return True for xl/worksheets/sheet*.xml entries."""
    return fname.startswith("xl/worksheets/") and fname.endswith(".xml")


def _apply_mapping_to_xml(text: str, mapping: dict[str, str]) -> str:
    """Replace keywords in raw XML text.

    Handles both plain form and XML-encoded form (e.g., '&' → '&amp;')
    so keywords containing special XML characters are matched correctly.
    """
    for kw, token in mapping.items():
        text = text.replace(kw, token)
        kw_enc = html.escape(kw, quote=False)
        if kw_enc != kw:
            text = text.replace(kw_enc, html.escape(token, quote=False))
    return text


def safe_local_defined_names(xlsx_path: Path) -> list[tuple[str, str, int | None]]:
    """Extract safe, local defined names directly from workbook.xml.

    openpyxl's DefinedNameDict is keyed by name only, so when the same name
    appears twice (once sheet-scoped with localSheetId, once as a broken
    workbook-level external ref), the external #REF! entry silently overwrites
    the valid local one.  We read the raw XML first to capture local definitions
    that openpyxl would lose.

    Returns list of (name, attr_text, local_sheet_id) for names that:
    - reference cells inside this workbook (no [n] external-workbook notation)
    - are not broken (#REF!)
    - are not _xlnm.* built-ins
    - are not hidden
    """
    with zipfile.ZipFile(xlsx_path) as z:
        with z.open("xl/workbook.xml") as f:
            tree = ET.parse(f)

    results = []
    for dn in tree.iter(f"{{{_XLSX_NS}}}definedName"):
        name = dn.get("name", "")
        value = dn.text or ""
        hidden = dn.get("hidden", "0") == "1"
        local_sheet_id_raw = dn.get("localSheetId")

        if name.startswith("_xlnm."):
            continue
        if hidden:
            continue
        if "#REF" in value or re.search(r"\[\d+\]", value):
            continue

        sid = int(local_sheet_id_raw) if local_sheet_id_raw is not None else None
        results.append((name, value, sid))

    return results


def extract_metadata(xlsx_path: Path) -> dict[str, str]:
    """Extract PII-containing metadata fields from an xlsx file.

    Reads both docProps/core.xml (author, title, subject, …) and
    docProps/app.xml (company, application name) directly from the zip.

    Args:
        xlsx_path: Path to the xlsx file.

    Returns:
        Dict mapping field names to non-empty string values.
    """
    meta: dict[str, str] = {}

    with zipfile.ZipFile(xlsx_path) as z:
        names = z.namelist()

        # --- core.xml ---
        if "docProps/core.xml" in names:
            with z.open("docProps/core.xml") as f:
                tree = ET.parse(f)
            root = tree.getroot()

            def _core(tag_ns: str, tag_local: str) -> str:
                el = root.find(f"{{{tag_ns}}}{tag_local}")
                return (el.text or "").strip() if el is not None else ""

            _fields = [
                ("creator", _NS_DC, "creator"),
                ("lastModifiedBy", _NS_CP, "lastModifiedBy"),
                ("title", _NS_DC, "title"),
                ("subject", _NS_DC, "subject"),
                ("description", _NS_DC, "description"),
                ("keywords", _NS_CP, "keywords"),
            ]
            for key, ns, local in _fields:
                val = _core(ns, local)
                if val:
                    meta[key] = val

        # --- app.xml ---
        if "docProps/app.xml" in names:
            with z.open("docProps/app.xml") as f:
                tree = ET.parse(f)
            root = tree.getroot()

            for local in ("Company", "Application"):
                el = root.find(f"{{{_NS_APP}}}{local}")
                val = (el.text or "").strip() if el is not None else ""
                if val:
                    meta[local.lower()] = val

    logger.info("Extracted metadata from %s: %s", xlsx_path.name, list(meta.keys()))
    return meta


def anonymize_xlsx(
    source_path: Path,
    dest_path: Path,
    mapping: dict[str, str],
) -> dict[str, str]:
    """Replace keywords in xlsx content and strip file metadata.

    Uses direct ZIP/XML manipulation instead of loading the workbook with
    openpyxl, keeping memory at O(single XML file) rather than O(full workbook).

    Keyword replacement targets xl/sharedStrings.xml (shared string table) and
    all xl/worksheets/*.xml (inline strings). Metadata is stripped by replacing
    docProps/core.xml and docProps/app.xml with clean versions. All other content
    (styles, drawings, defined names, charts, etc.) is copied byte-for-byte so
    nothing is lost.

    Args:
        source_path: Path to the source xlsx file (working.xlsx).
        dest_path: Path to save the anonymized xlsx file.
        mapping: Dict of {original_keyword: replacement} pairs.

    Returns:
        The mapping dict, sorted by key length descending (replacement order).
    """
    mapping = dict(
        sorted(
            {k.strip(): v.strip() for k, v in mapping.items() if k.strip()}.items(),
            key=lambda item: len(item[0]),
            reverse=True,
        )
    )

    replaced_parts: list[str] = []
    tmp_path = dest_path.with_suffix(".tmp.xlsx")
    try:
        with zipfile.ZipFile(source_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    fname = item.filename
                    if fname == "docProps/core.xml":
                        zout.writestr(item, _CLEAN_CORE_XML)
                    elif fname == "docProps/app.xml":
                        zout.writestr(item, _CLEAN_APP_XML)
                    elif fname in _XLSX_TEXT_TARGETS or _is_worksheet_xml(fname):
                        text = zin.read(fname).decode("utf-8")
                        patched = _apply_mapping_to_xml(text, mapping)
                        if patched != text:
                            replaced_parts.append(fname)
                        zout.writestr(item, patched.encode("utf-8"))
                    else:
                        zout.writestr(item, zin.read(fname))
        shutil.move(str(tmp_path), str(dest_path))
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise

    logger.info(
        "Anonymized %s -> %s: %d keywords, replacements in: %s",
        source_path.name,
        dest_path.name,
        len(mapping),
        replaced_parts or "none",
    )
    return mapping


def extract_metadata_docx(docx_path: Path) -> dict[str, str]:
    """Extract PII-containing metadata fields from a docx file.

    Args:
        docx_path: Path to the docx file.

    Returns:
        Dict mapping field names to non-empty string values.
    """
    doc = DocxDocument(docx_path)
    cp = doc.core_properties
    meta: dict[str, str] = {}
    for key, val in [
        ("creator", cp.author),
        ("lastModifiedBy", cp.last_modified_by),
        ("title", cp.title),
        ("subject", cp.subject),
        ("keywords", cp.keywords),
    ]:
        if val and str(val).strip():
            meta[key] = str(val).strip()
    logger.info("Extracted docx metadata from %s: %s", docx_path.name, list(meta.keys()))
    return meta


def _strip_docx_layout(doc: "DocxDocument") -> None:  # type: ignore[name-defined]
    """Remove headers, footers, and background from a docx Document (in-place).

    - Headers/footers: clears all runs in every section's header and footer
      objects (default, first-page, even-page variants).
    - Background: removes the <w:background> element from the document root,
      which carries the page background color or fill.
    """
    from docx.oxml.ns import qn  # local import — only needed here

    # --- headers & footers ---
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                if not hf.is_linked_to_previous:
                    for para in hf.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ""
            except Exception:
                pass  # some hf objects raise if the part doesn't exist

    # --- background ---
    bg = doc.element.find(qn("w:background"))
    if bg is not None:
        doc.element.remove(bg)


def anonymize_docx(
    source_path: Path,
    dest_path: Path,
    mapping: dict[str, str],
    strip_layout: bool = False,
) -> dict[str, str]:
    """Replace keywords in all docx text (paragraphs + tables) and strip metadata.

    Keywords are sorted by length descending to avoid partial replacements.

    Args:
        source_path: Path to the source docx file.
        dest_path: Path to save the anonymized docx file.
        mapping: Dict of {original_keyword: replacement} pairs.
        strip_layout: If True, also remove headers, footers, and page background.
            Should be set for policy/politiques documents only.

    Returns:
        The mapping dict, sorted by key length descending.
    """
    mapping = dict(
        sorted(
            {k.strip(): v.strip() for k, v in mapping.items() if k.strip()}.items(),
            key=lambda item: len(item[0]),
            reverse=True,
        )
    )

    shutil.copy2(source_path, dest_path)
    doc = DocxDocument(dest_path)

    def _replace_in_paragraph(para) -> int:
        count = 0
        for run in para.runs:
            if run.text:
                original = run.text
                for kw, token in mapping.items():
                    run.text = run.text.replace(kw, token)
                if run.text != original:
                    count += 1
        return count

    replaced_count = 0
    for para in doc.paragraphs:
        replaced_count += _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replaced_count += _replace_in_paragraph(para)

    if strip_layout:
        _strip_docx_layout(doc)

    cp = doc.core_properties
    cp.author = "ANONYME"
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.company = ""

    doc.save(dest_path)

    logger.info(
        "Anonymized docx %s -> %s: %d keywords, %d runs modified%s",
        source_path.name, dest_path.name, len(mapping), replaced_count,
        " (layout stripped)" if strip_layout else "",
    )
    return mapping


def _apply_catamania_fix_xlsx(xlsx_path: Path) -> None:
    """Replace 'Cat-Amania' (case-insensitive) with 'Catamania' in xlsx text content (in-place).

    Applied as a second pass after de-anonymization so that any occurrence
    introduced by Claude (despite instructions to use FOURNISSEUR) is corrected.
    """
    tmp_path = xlsx_path.with_suffix(".tmp.xlsx")
    try:
        with zipfile.ZipFile(xlsx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    fname = item.filename
                    if fname in _XLSX_TEXT_TARGETS or _is_worksheet_xml(fname):
                        text = zin.read(fname).decode("utf-8")
                        patched = _CATAMANIA_RE.sub("Catamania", text)
                        zout.writestr(item, patched.encode("utf-8"))
                    else:
                        zout.writestr(item, zin.read(fname))
        shutil.move(str(tmp_path), str(xlsx_path))
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise


def deanonymize_xlsx(
    source_path: Path,
    dest_path: Path,
    mapping: dict[str, str],
) -> None:
    """Reverse the anonymization of an xlsx file.

    Replaces tokens with their original values by inverting the mapping and
    calling anonymize_xlsx (which already handles length-sorted replacement).
    Then applies the Catamania fix (case-insensitive 'Cat-Amania' → 'Catamania').

    Args:
        source_path: Path to the anonymized xlsx (output_anon.xlsx).
        dest_path: Path to save the de-anonymized xlsx (output.xlsx).
        mapping: The original anonymization mapping {original: token} as stored
                 in anonymized_map.json.
    """
    inverse = {token: original for original, token in mapping.items() if token}
    anonymize_xlsx(source_path, dest_path, inverse)
    _apply_catamania_fix_xlsx(dest_path)
    logger.info("De-anonymized %s -> %s", source_path.name, dest_path.name)


def deanonymize_text(text: str, mapping: dict[str, str]) -> str:
    """Replace anonymization tokens with original values in a plain-text string.

    Sorts by original value length descending to avoid partial replacements.

    Args:
        text: Text that may contain anonymization tokens.
        mapping: The original anonymization mapping {original: token}.

    Returns:
        Text with tokens replaced by original values.
    """
    # Sort by original (value) length descending to prevent partial matches
    for original, token in sorted(mapping.items(), key=lambda x: len(x[1]), reverse=True):
        if token:
            text = text.replace(token, original)
    # Second pass: fix 'Cat-Amania' → 'Catamania' (case-insensitive)
    text = _CATAMANIA_RE.sub("Catamania", text)
    return text


def _strip_app_xml(xlsx_path: Path) -> None:
    """Replace docProps/app.xml with a minimal clean version (in-place).

    Rewrites the xlsx zip, substituting a clean app.xml that contains no
    company name, application name, or other extended properties.
    """
    tmp_path = xlsx_path.with_suffix(".tmp.xlsx")
    try:
        with zipfile.ZipFile(xlsx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "docProps/app.xml":
                        zout.writestr(item, _CLEAN_APP_XML)
                    else:
                        zout.writestr(item, zin.read(item.filename))
        shutil.move(str(tmp_path), str(xlsx_path))
    except Exception:
        tmp_path.unlink(missing_ok=True)
        raise
